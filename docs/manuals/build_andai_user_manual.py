from __future__ import annotations

from dataclasses import dataclass
from html import escape
from pathlib import Path
import re

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL_DIR = ROOT / "docs" / "manuals"
SCREENSHOT_DIR = ROOT / "docs" / "assets" / "manual_screenshots"
PDF_PATH = MANUAL_DIR / "ANDAI_User_Manual.pdf"
MD_PATH = MANUAL_DIR / "ANDAI_User_Manual.md"
DOCX_PATH = MANUAL_DIR / "ANDAI_User_Manual.docx"
LOGO_PATH = ROOT / "docs" / "assets" / "branding" / "andai_logo_transparent.png"

TITLE = "ANDAI Platform User Manual"
SUBTITLE = "Step-by-step training guide for normal users, organization admins, and full admins"
VERSION = "Version 4.0"
DATE = "Prepared on: July 14, 2026"
AUDIENCE = "Audience: ANDAI platform users and administrators"


@dataclass
class Section:
    title: str
    blocks: list[dict]


SECTIONS = [
    Section(
        "1. Purpose and Scope",
        [
            {
                "type": "p",
                "text": (
                    "This manual teaches Full Admins, Organization Admins, and Normal Users how to use ANDAI. "
                    "Follow only the section for your role; ANDAI hides pages that your account is not allowed to use."
                ),
            },
            {
                "type": "callout",
                "title": "How ANDAI protects knowledge",
                "text": (
                    "Your role controls the pages you can open. Your organization and department access control which documents, FAQ items, database tables, and chat evidence you can use."
                ),
            },
            {
                "type": "list",
                "title": "What this manual covers",
                "items": [
                    "Signing in and finding the pages available to you.",
                    "Asking questions and checking the evidence used in an answer.",
                    "Uploading and maintaining approved organization knowledge.",
                    "Creating organizations, departments, and users.",
                    "Testing answer quality, reviewing usage, and checking audit logs.",
                    "Using the authenticated chat API for authorized technical testing.",
                ],
            },
        ],
    ),
    Section(
        "2. Roles and Access",
        [
            {
                "type": "table",
                "headers": ["Role", "Pages shown", "Main responsibility", "Limits"],
                "rows": [
                    ["Normal User", "Assistant", "Ask questions and verify answers.", "Cannot upload knowledge or manage accounts."],
                    ["Organization Admin", "Overview, Assistant, Documents, FAQ, Database, Evaluations, Analytics, Users, Audit Logs", "Maintain one organization's knowledge and users.", "Limited to the assigned organization and departments."],
                    ["Full Admin / Super Admin", "Overview, Organizations, Users, Audit Logs", "Set up and oversee the platform.", "Daily knowledge maintenance belongs to Organization Admins."],
                ],
            },
            {
                "type": "callout",
                "title": "Department access",
                "text": (
                    "A user without a department can still open Assistant and chat. However, private organization knowledge may return no evidence until an admin assigns a department. New organizations receive a General department by default."
                ),
            },
        ],
    ),
    Section(
        "3. Sign In and Navigate",
        [
            {
                "type": "steps",
                "title": "Sign in",
                "items": [
                    "Open https://andai.my or the address supplied by your administrator.",
                    "Enter the email address and temporary password supplied to you.",
                    "Click Sign in.",
                    "Confirm that the name and role shown at the bottom of the sidebar are correct.",
                ],
            },
            {
                "type": "image",
                "path": "01-login.png",
                "caption": "Screenshot 1: ANDAI sign-in page.",
            },
            {
                "type": "list",
                "title": "Main screen areas",
                "items": [
                    "Sidebar: opens the pages allowed for your role.",
                    "Main area: shows the page you selected.",
                    "Account area: shows your name, email, role, and Sign out.",
                    "Organization selector: shown only when your role can work across organizations.",
                ],
            },
            {
                "type": "callout",
                "title": "A missing page is usually not an error",
                "text": "ANDAI deliberately hides pages that do not belong to your role. Ask an administrator to check your role if you believe access is missing.",
            },
        ],
    ),
    Section(
        "4. Normal User - Ask ANDAI",
        [
            {
                "type": "steps",
                "title": "Start a conversation",
                "items": [
                    "Open Assistant.",
                    "Click New Chat when you want to start a separate topic.",
                    "Keep Database, PDF / Docs, and FAQ selected unless you want to search fewer sources.",
                    "Type one clear question in the message box.",
                    "Click the send button and wait for the answer.",
                    "Open an earlier conversation from the list on the left when you need to continue it.",
                ],
            },
            {
                "type": "image",
                "path": "15-normal-user-assistant.png",
                "caption": "Screenshot 2: Normal User view. Only Assistant and chat history are available.",
            },
            {
                "type": "table",
                "headers": ["Source", "Use it for", "Example"],
                "rows": [
                    ["Database", "Numbers, lists, and facts from imported tables.", "How many applications were approved this month?"],
                    ["PDF / Docs", "Policies, manuals, procedures, and Word documents.", "What documents are required for this application?"],
                    ["FAQ", "Short, approved answers to common questions.", "Who should I contact for account support?"],
                ],
            },
            {
                "type": "list",
                "title": "Check every answer",
                "items": [
                    "Read the answer and the source badges or citations shown with it.",
                    "If ANDAI says it cannot find evidence, make the question more specific or ask an admin to add the missing knowledge.",
                    "Do not treat an answer as policy when no source is shown.",
                    "Start a new chat when changing to an unrelated topic.",
                ],
            },
        ],
    ),
    Section(
        "5. Organization Admin - Overview",
        [
            {
                "type": "p",
                "text": (
                    "Organization Admins maintain the knowledge used by people in their organization. "
                    "The account may be shown as ADMIN in the sidebar. Select the correct department before creating or uploading content."
                ),
            },
            {
                "type": "image",
                "path": "06-org-admin-overview.png",
                "caption": "Screenshot 3: Organization Admin overview and administration navigation.",
            },
            {
                "type": "table",
                "headers": ["Page", "Purpose"],
                "rows": [
                    ["Overview", "Check knowledge counts, quick actions, and service status."],
                    ["Assistant", "Test the same question experience used by normal users."],
                    ["Documents", "Upload PDF and Word documents."],
                    ["FAQ", "Create approved question-and-answer entries."],
                    ["Database", "Create tables and import CSV or SQL data."],
                    ["Evaluations", "Run repeatable answer-quality tests."],
                    ["Analytics", "Review usage and unanswered questions."],
                    ["Users", "Manage users and department access in the organization."],
                    ["Audit Logs", "Review important actions in the organization."],
                ],
            },
        ],
    ),
    Section(
        "6. Organization Admin - Documents and FAQ",
        [
            {
                "type": "steps",
                "title": "Upload a document",
                "items": [
                    "Open Documents.",
                    "Select the department that owns the document.",
                    "Click Upload Document and choose a readable PDF or Word .docx file.",
                    "Wait until the status becomes Ready before testing questions.",
                    "If processing fails, use Retry processing or upload a text-readable version.",
                ],
            },
            {
                "type": "image",
                "path": "07-org-admin-documents.png",
                "caption": "Screenshot 4: Organization Admin Documents page.",
            },
            {
                "type": "steps",
                "title": "Create an FAQ item",
                "items": [
                    "Open FAQ and select the correct department.",
                    "Click Add FAQ.",
                    "Enter the question in the words users normally use.",
                    "Enter the short, approved answer.",
                    "Publish and save the item when it is ready for users.",
                ],
            },
            {
                "type": "image",
                "path": "08-org-admin-faq.png",
                "caption": "Screenshot 5: Organization Admin FAQ page.",
            },
            {
                "type": "callout",
                "title": "Knowledge quality",
                "text": "Use final approved files, clear filenames, and short FAQ answers. Remove or unpublish outdated information when a policy changes.",
            },
        ],
    ),
    Section(
        "7. Organization Admin - Database",
        [
            {
                "type": "list",
                "title": "Available database actions",
                "items": [
                    "Tables: review imported tables and open their rows.",
                    "Create Table: define a small table and its columns manually.",
                    "Upload Table & Data: create a table from a CSV or SQL file.",
                    "Upload Data: append to or replace rows in an existing table.",
                ],
            },
            {
                "type": "image",
                "path": "09-org-admin-database.png",
                "caption": "Screenshot 6: Database page with table-management actions.",
            },
            {
                "type": "steps",
                "title": "Upload a new table",
                "items": [
                    "Open Database and select the owning department.",
                    "Click Upload Table & Data.",
                    "Enter a clear display name and optional description.",
                    "Choose a CSV or SQL file.",
                    "Review the preview when shown, then click Create & Import.",
                    "Return to Tables and use View data to confirm the imported rows.",
                ],
            },
            {
                "type": "image",
                "path": "10-org-admin-database-upload.png",
                "caption": "Screenshot 7: Upload Table & Data form.",
            },
        ],
    ),
    Section(
        "8. Organization Admin - Evaluations and Analytics",
        [
            {
                "type": "steps",
                "title": "Test answer quality",
                "items": [
                    "Open Evaluations.",
                    "Add a representative question.",
                    "Enter expected keywords and an expected source when useful.",
                    "Run the test and review whether it passed, its answer, sources, and response time.",
                    "Improve the underlying document, FAQ, or table when the answer is weak.",
                ],
            },
            {
                "type": "image",
                "path": "11-org-admin-evaluations.png",
                "caption": "Screenshot 8: Evaluation Tests page.",
            },
            {
                "type": "list",
                "title": "Use analytics to improve knowledge",
                "items": [
                    "Review chat volume and active users.",
                    "Check which sources are being used.",
                    "Review common and unanswered questions.",
                    "Add or improve knowledge for questions users cannot answer.",
                ],
            },
            {
                "type": "image",
                "path": "12-org-admin-analytics.png",
                "caption": "Screenshot 9: Usage Analytics page.",
            },
        ],
    ),
    Section(
        "9. Organization Admin - Users and Audit Logs",
        [
            {
                "type": "steps",
                "title": "Create or update a user",
                "items": [
                    "Open Users.",
                    "Click Add User to create an account in your organization.",
                    "Enter the user's name, email, temporary password, and role.",
                    "Assign General or the departments the user is allowed to access.",
                    "Use Department Access later when the user's responsibilities change.",
                ],
            },
            {
                "type": "image",
                "path": "13-org-admin-users.png",
                "caption": "Screenshot 10: Organization Admin Users page.",
            },
            {
                "type": "list",
                "title": "Review audit activity",
                "items": [
                    "Open Audit Logs to review logins, uploads, user changes, and other important actions.",
                    "Use the displayed user and organization names to understand who performed each action.",
                    "Record the time and details before escalating unusual activity to the Full Admin.",
                ],
            },
            {
                "type": "image",
                "path": "14-org-admin-audit-logs.png",
                "caption": "Screenshot 11: Organization Admin Audit Logs page.",
            },
        ],
    ),
    Section(
        "10. Full Admin - Platform Setup",
        [
            {
                "type": "p",
                "text": "Full Admins manage platform structure and access. Their sidebar intentionally excludes daily Assistant and knowledge-maintenance pages.",
            },
            {
                "type": "image",
                "path": "02-full-admin-overview.png",
                "caption": "Screenshot 12: Full Admin overview.",
            },
            {
                "type": "steps",
                "title": "Create an organization and departments",
                "items": [
                    "Open Organizations and click Add Organization.",
                    "Enter the organization name and description.",
                    "Create the organization; ANDAI adds a General department automatically.",
                    "Open Manage Departments to add units such as HR, Finance, Operations, or Credit.",
                    "Keep departments separate when their knowledge must not be shared.",
                ],
            },
            {
                "type": "image",
                "path": "03-full-admin-organizations.png",
                "caption": "Screenshot 13: Organizations and department controls.",
            },
            {
                "type": "steps",
                "title": "Create users and assign roles",
                "items": [
                    "Open Users and click Add User.",
                    "Enter the person's name, email, temporary password, role, and organization.",
                    "Assign department access for Organization Admins and Normal Users.",
                    "Create the account and verify its role in the user list.",
                    "Give the temporary password to the user through a separate secure channel.",
                ],
            },
            {
                "type": "image",
                "path": "04-full-admin-users.png",
                "caption": "Screenshot 14: Full Admin Users page.",
            },
            {
                "type": "image",
                "path": "05-full-admin-audit-logs.png",
                "caption": "Screenshot 15: Full Admin Audit Logs across organizations.",
            },
        ],
    ),
    Section(
        "11. Troubleshooting",
        [
            {
                "type": "table",
                "headers": ["Problem", "Likely cause", "What to do"],
                "rows": [
                    ["I only see Assistant.", "The account is a Normal User.", "This is expected. Ask an admin if the role is wrong."],
                    ["I can chat but get no organization evidence.", "No department access or no matching knowledge.", "Ask an admin to assign a department and confirm that knowledge is Ready or published."],
                    ["ANDAI cannot find the answer.", "The question is broad or the selected sources contain no match.", "Ask more specifically, select the relevant sources, or add the missing knowledge."],
                    ["A document stays in Processing or Error.", "The file is scanned, unreadable, or processing failed.", "Use OCR for scanned files, retry processing, or upload a text-readable PDF or DOCX."],
                    ["Upload or Save appears to do nothing.", "A required field, file, or department is missing.", "Check every required field and confirm a department is selected."],
                    ["A database answer looks wrong.", "The table is outdated or owned by the wrong department.", "Use View data, confirm the rows and department, then re-import if needed."],
                    ["A page is missing.", "The role does not include that page.", "Ask an administrator to check the user's role and organization."],
                ],
            },
            {
                "type": "callout",
                "title": "Information to include in a support request",
                "text": "Provide your email, role, organization, department, page name, what you clicked, the exact question or filename, and the time the problem occurred. Never send your password or access token.",
            },
        ],
    ),
    Section(
        "12. Suggested Training Session",
        [
            {
                "type": "steps",
                "title": "45-minute training flow",
                "items": [
                    "5 minutes: explain roles, organizations, and department access.",
                    "10 minutes: Normal User starts a chat, selects sources, and checks citations.",
                    "10 minutes: Organization Admin uploads one document and creates one FAQ item.",
                    "8 minutes: Organization Admin imports a small CSV and confirms its rows.",
                    "7 minutes: Full Admin shows organization, department, and user setup.",
                    "5 minutes: review troubleshooting and the support process.",
                ],
            },
            {
                "type": "list",
                "title": "Trainer preparation",
                "items": [
                    "Use non-sensitive demo data and one test account for each role.",
                    "Prepare one readable PDF or DOCX, one approved FAQ, and one small CSV.",
                    "Confirm the Organization Admin has department access before training.",
                    "Confirm the document is Ready before asking questions about it.",
                    "Share temporary passwords separately; never place them in training materials.",
                ],
            },
        ],
    ),
    Section(
        "Appendix A. API Testing for Technical Users",
        [
            {
                "type": "p",
                "text": (
                    "Authorized technical users can ask ANDAI directly through the authenticated API. "
                    "This is separate from the standard web controls and follows the same user, organization, and department permissions."
                ),
            },
            {
                "type": "steps",
                "title": "Basic API test",
                "items": [
                    "Request an access token from POST /api/auth/login using your assigned email and password.",
                    "Copy only the access_token value from the response.",
                    "Send a JSON question to POST /api/chat with the token in the Authorization header.",
                    "Check the answer and sources in the JSON response.",
                ],
            },
            {
                "type": "code",
                "language": "bash",
                "text": (
                    "curl -X POST https://andai.my/api/auth/login \\\n"
                    "  -H \"Content-Type: application/x-www-form-urlencoded\" \\\n"
                    "  --data-urlencode \"username=YOUR_EMAIL\" \\\n"
                    "  --data-urlencode \"password=YOUR_PASSWORD\""
                ),
            },
            {
                "type": "code",
                "language": "bash",
                "text": (
                    "curl -X POST https://andai.my/api/chat \\\n"
                    "  -H \"Authorization: Bearer YOUR_ACCESS_TOKEN\" \\\n"
                    "  -H \"Content-Type: application/json\" \\\n"
                    "  -d '{\"message\":\"What is the approved procedure?\",\"sources\":[\"documents\",\"faq\"],\"ai_insights\":false}'"
                ),
            },
            {
                "type": "table",
                "headers": ["Response", "Meaning", "Action"],
                "rows": [
                    ["200", "The request was accepted.", "Read message and sources in the JSON response."],
                    ["401", "The token is missing, invalid, or expired.", "Sign in again and use the new access token."],
                    ["403", "The account cannot access the requested organization or department.", "Ask an admin to check the account assignment."],
                    ["No evidence", "Approved sources did not contain a match.", "Ask more specifically or improve the approved knowledge."],
                ],
            },
            {
                "type": "callout",
                "title": "Protect credentials",
                "text": "Do not paste passwords or access tokens into tickets, screenshots, documents, or chat messages. Use a test account and non-sensitive questions during demonstrations.",
            },
        ],
    ),
]


def inline(text: str) -> str:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    rendered = []
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            rendered.append(f"<b>{escape(part[2:-2])}</b>")
        else:
            rendered.append(escape(part))
    return "".join(rendered)


def markdown() -> str:
    lines = [
        f"# {TITLE}",
        "",
        SUBTITLE,
        "",
        f"- {VERSION}",
        f"- {DATE}",
        f"- {AUDIENCE}",
        "",
        "## Contents",
        "",
    ]
    for section in SECTIONS:
        lines.append(f"- {section.title}")
    lines.append("")

    for section in SECTIONS:
        lines.extend([f"## {section.title}", ""])
        for block in section.blocks:
            block_type = block["type"]
            if block_type == "p":
                lines.extend([block["text"], ""])
            elif block_type == "callout":
                lines.extend([f"**{block['title']}:** {block['text']}", ""])
            elif block_type in {"list", "steps"}:
                lines.extend([f"### {block['title']}", ""])
                marker_items = enumerate(block["items"], start=1) if block_type == "steps" else block["items"]
                if block_type == "steps":
                    for index, item in marker_items:
                        lines.append(f"{index}. {item}")
                else:
                    for item in marker_items:
                        lines.append(f"- {item}")
                lines.append("")
            elif block_type == "table":
                headers = block["headers"]
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row in block["rows"]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
            elif block_type == "image":
                lines.append(f"![{block['caption']}](../assets/manual_screenshots/{block['path']})")
                lines.append("")
                lines.append(f"*{block['caption']}*")
                lines.append("")
            elif block_type == "code":
                lines.extend([f"```{block.get('language', '')}", block["text"], "```", ""])
    return "\n".join(lines).rstrip() + "\n"


def build_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ManualTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=28,
            leading=32,
            textColor=colors.HexColor("#111827"),
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "ManualSubtitle",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=12,
            leading=17,
            textColor=colors.HexColor("#475569"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "h1": ParagraphStyle(
            "ManualH1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#B91C1C"),
            spaceBefore=12,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "ManualH2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#1F2937"),
            spaceBefore=8,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "ManualBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "ManualSmall",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#475569"),
        ),
        "bullet": ParagraphStyle(
            "ManualBullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=12.8,
            leftIndent=14,
            firstLineIndent=-8,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=4,
        ),
        "toc": ParagraphStyle(
            "ManualTOC",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=5,
        ),
        "table_header": ParagraphStyle(
            "ManualTableHeader",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8,
            leading=10,
            textColor=colors.white,
        ),
        "table_cell": ParagraphStyle(
            "ManualTableCell",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.6,
            leading=9.4,
            textColor=colors.HexColor("#1F2937"),
        ),
        "callout_title": ParagraphStyle(
            "ManualCalloutTitle",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#7F1D1D"),
        ),
        "code": ParagraphStyle(
            "ManualCode",
            parent=base["Code"],
            fontName="Courier",
            fontSize=7.2,
            leading=9,
            textColor=colors.HexColor("#111827"),
        ),
    }


def page_footer(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#E5E7EB"))
    canvas.line(doc.leftMargin, 0.58 * inch, LETTER[0] - doc.rightMargin, 0.58 * inch)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#64748B"))
    canvas.drawString(doc.leftMargin, 0.38 * inch, "ANDAI Platform User Manual")
    canvas.drawRightString(LETTER[0] - doc.rightMargin, 0.38 * inch, f"Page {doc.page}")
    canvas.restoreState()


def table_flow(headers: list[str], rows: list[list[str]], styles, col_widths=None):
    width = 7.0 * inch
    if col_widths is None:
        col_widths = [width / len(headers)] * len(headers)
    data = [[Paragraph(inline(h), styles["table_header"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(inline(cell), styles["table_cell"]) for cell in row])
    table = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#B91C1C")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ]
        )
    )
    return table


def callout_flow(title: str, text: str, styles):
    content = [
        [Paragraph(inline(title), styles["callout_title"])],
        [Paragraph(inline(text), styles["body"])],
    ]
    table = Table(content, colWidths=[7.0 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FEF2F2")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#FCA5A5")),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return KeepTogether([table, Spacer(1, 8)])


def image_flow(image_name: str, caption: str, styles):
    image_path = SCREENSHOT_DIR / image_name
    if not image_path.exists():
        raise FileNotFoundError(f"Missing manual screenshot: {image_path}")

    image_reader = ImageReader(str(image_path))
    image_width, image_height = image_reader.getSize()
    max_width = 6.8 * inch
    max_height = 4.7 * inch
    width = max_width
    height = width * image_height / image_width
    if height > max_height:
        height = max_height
        width = height * image_width / image_height

    image = Image(str(image_path), width=width, height=height)
    image.hAlign = "CENTER"
    return [
        image,
        Spacer(1, 4),
        Paragraph(inline(caption), styles["small"]),
        Spacer(1, 10),
    ]


def code_flow(text: str, styles):
    code_html = "<br/>".join(escape(line).replace(" ", "&#160;") for line in text.splitlines())
    table = Table([[Paragraph(code_html, styles["code"])]], colWidths=[7.0 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return [table, Spacer(1, 8)]


def add_list(story, title: str, items: list[str], ordered: bool, styles):
    story.append(Paragraph(inline(title), styles["h2"]))
    for index, item in enumerate(items, start=1):
        prefix = f"{index}. " if ordered else "- "
        story.append(Paragraph(prefix + inline(item), styles["bullet"]))
    story.append(Spacer(1, 5))


def build_pdf():
    styles = build_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
        title=TITLE,
        author="ANDAI",
        subject="User manual",
    )
    story = []

    if LOGO_PATH.exists():
        logo = Image(str(LOGO_PATH), width=1.15 * inch, height=1.15 * inch)
        logo.hAlign = "CENTER"
        story.extend([Spacer(1, 0.55 * inch), logo, Spacer(1, 20)])
    else:
        story.append(Spacer(1, 1.1 * inch))

    story.append(Paragraph(TITLE, styles["title"]))
    story.append(Paragraph(SUBTITLE, styles["subtitle"]))
    story.append(Spacer(1, 18))
    story.append(Paragraph(VERSION, styles["subtitle"]))
    story.append(Paragraph(DATE, styles["subtitle"]))
    story.append(Paragraph(AUDIENCE, styles["subtitle"]))
    story.append(Spacer(1, 28))
    story.append(
        table_flow(
            ["Role", "Core learning outcome"],
            [
                ["Normal User", "Ask questions and verify answer sources."],
                ["Organization Admin", "Maintain knowledge sources and department access."],
                ["Full Admin", "Set up organizations, departments, users, and oversight."],
            ],
            styles,
            [1.45 * inch, 5.55 * inch],
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("Contents", styles["h1"]))
    for section in SECTIONS:
        story.append(Paragraph(section.title, styles["toc"]))
    story.append(PageBreak())

    for section_index, section in enumerate(SECTIONS):
        if section_index and section.title.startswith("Appendix"):
            story.append(PageBreak())
        story.append(Paragraph(section.title, styles["h1"]))
        for block in section.blocks:
            block_type = block["type"]
            if block_type == "p":
                story.append(Paragraph(inline(block["text"]), styles["body"]))
            elif block_type == "callout":
                story.append(callout_flow(block["title"], block["text"], styles))
            elif block_type == "list":
                add_list(story, block["title"], block["items"], False, styles)
            elif block_type == "steps":
                add_list(story, block["title"], block["items"], True, styles)
            elif block_type == "table":
                headers = block["headers"]
                if len(headers) == 4:
                    widths = [1.05 * inch, 1.9 * inch, 2.0 * inch, 2.05 * inch]
                elif len(headers) == 3:
                    widths = [1.55 * inch, 2.55 * inch, 2.9 * inch]
                else:
                    widths = None
                story.append(table_flow(headers, block["rows"], styles, widths))
                story.append(Spacer(1, 10))
            elif block_type == "image":
                story.extend(image_flow(block["path"], block["caption"], styles))
            elif block_type == "code":
                story.extend(code_flow(block["text"], styles))
        if section_index < len(SECTIONS) - 1:
            story.append(Spacer(1, 7))

    doc.build(story, onFirstPage=page_footer, onLaterPages=page_footer)


def set_docx_run(run, size=None, color=None, bold=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def docx_paragraph(doc, text: str, style: str | None = None, bold=False, color=None, size=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_docx_run(run, size=size, color=color, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    return paragraph


def docx_list_paragraph(doc, text: str, prefix: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.1)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(prefix + text)
    set_docx_run(run, size=9.2, color=RGBColor(31, 41, 55))
    return paragraph


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, header=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_docx_run(
        run,
        size=8.5 if header else 8,
        color=RGBColor(255, 255, 255) if header else RGBColor(31, 41, 55),
        bold=header,
    )
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def docx_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, header=True)
        set_cell_shading(header_cells[index], "B91C1C")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    doc.add_paragraph()
    return table


def docx_callout(doc, title: str, text: str):
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    for cell in table.columns[0].cells:
        set_cell_shading(cell, "FEF2F2")
    set_cell_text(table.cell(0, 0), title)
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(127, 29, 29)
    set_cell_text(table.cell(1, 0), text)
    doc.add_paragraph()


def docx_image(doc, image_name: str, caption: str):
    image_path = SCREENSHOT_DIR / image_name
    if not image_path.exists():
        raise FileNotFoundError(f"Missing manual screenshot: {image_path}")

    image_width, image_height = ImageReader(str(image_path)).getSize()
    width_inches = 6.35
    height_inches = width_inches * image_height / image_width
    max_height_inches = 4.7
    if height_inches > max_height_inches:
        height_inches = max_height_inches
        width_inches = height_inches * image_width / image_height

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches), height=Inches(height_inches))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(caption)
    set_docx_run(run, size=8, color=RGBColor(71, 85, 105))


def docx_code(doc, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(17, 24, 39)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(185, 28, 28)),
        ("Heading 2", 12, RGBColor(31, 41, 55)),
        ("Heading 3", 11, RGBColor(31, 41, 55)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    footer = section.footer.paragraphs[0]
    footer.text = "ANDAI Platform User Manual"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(0.7))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(TITLE)
    set_docx_run(title_run, size=24, color=RGBColor(17, 24, 39), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(SUBTITLE)
    set_docx_run(subtitle_run, size=11, color=RGBColor(71, 85, 105))

    for line in [VERSION, DATE, AUDIENCE]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_docx_run(run, size=10, color=RGBColor(71, 85, 105))

    docx_table(
        doc,
        ["Role", "Core learning outcome"],
        [
            ["Normal User", "Ask questions and verify answer sources."],
            ["Organization Admin", "Maintain knowledge sources and department access."],
            ["Full Admin", "Set up organizations, departments, users, and oversight."],
        ],
    )
    doc.add_page_break()

    doc.add_heading("Contents", level=1)
    for section_item in SECTIONS:
        docx_paragraph(doc, section_item.title)
    doc.add_page_break()

    for section_index, section_item in enumerate(SECTIONS):
        if section_index and section_item.title.startswith("Appendix"):
            doc.add_page_break()
        doc.add_heading(section_item.title, level=1)
        for block in section_item.blocks:
            block_type = block["type"]
            if block_type == "p":
                docx_paragraph(doc, block["text"])
            elif block_type == "callout":
                docx_callout(doc, block["title"], block["text"])
            elif block_type == "table":
                docx_table(doc, block["headers"], block["rows"])
            elif block_type == "image":
                docx_image(doc, block["path"], block["caption"])
            elif block_type == "code":
                docx_code(doc, block["text"])
            elif block_type in {"list", "steps"}:
                doc.add_heading(block["title"], level=2)
                for item_index, item in enumerate(block["items"], start=1):
                    prefix = f"{item_index}. " if block_type == "steps" else "- "
                    docx_list_paragraph(doc, item, prefix)

    doc.save(DOCX_PATH)


def main():
    MANUAL_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(markdown(), encoding="utf-8")
    build_pdf()
    build_docx()
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {PDF_PATH}")
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
