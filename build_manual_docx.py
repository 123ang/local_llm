from __future__ import annotations

import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/Users/123ang/andai-runtime/local_llm")
HTML_PATH = ROOT / "AskAI_Web_User_Manual.html"
DOCX_PATH = ROOT / "AskAI_Web_User_Manual.docx"


def normalized_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, code: bool = False):
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def append_inline(paragraph, node, *, bold: bool = False, italic: bool = False, code: bool = False):
    if isinstance(node, NavigableString):
        text = str(node)
        if not text.strip():
            if "\n" not in text:
                add_run(paragraph, " ", bold=bold, italic=italic, code=code)
            return
        add_run(paragraph, re.sub(r"\s+", " ", text), bold=bold, italic=italic, code=code)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run("\n")
        return

    next_bold = bold or node.name in {"strong", "b"}
    next_italic = italic or node.name in {"em", "i"}
    next_code = code or node.name == "code"

    for child in node.children:
        append_inline(paragraph, child, bold=next_bold, italic=next_italic, code=next_code)


def paragraph_from_tag(document: Document, tag: Tag, *, style: str | None = None, align=None):
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    for child in tag.children:
        append_inline(paragraph, child)
    return paragraph


def list_from_tag(document: Document, tag: Tag):
    style = "List Bullet" if tag.name == "ul" else "List Number"
    for li in tag.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style)
        for child in li.children:
            if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                nested_style = "List Bullet 2" if child.name == "ul" else "List Number 2"
                for nested_li in child.find_all("li", recursive=False):
                    nested_p = document.add_paragraph(style=nested_style)
                    for nested_child in nested_li.children:
                        append_inline(nested_p, nested_child)
            else:
                append_inline(paragraph, child)


def table_from_tag(document: Document, tag: Tag):
    rows = tag.find_all("tr")
    if not rows:
      return

    col_count = max(len(row.find_all(["th", "td"])) for row in rows)
    table = document.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"

    for row in rows:
        cells = row.find_all(["th", "td"])
        doc_row = table.add_row()
        for idx in range(col_count):
            cell = doc_row.cells[idx]
            cell.text = ""
            if idx < len(cells):
                paragraph = cell.paragraphs[0]
                for child in cells[idx].children:
                    append_inline(paragraph, child)
                if cells[idx].name == "th":
                    for run in paragraph.runs:
                        run.bold = True


def figure_from_tag(document: Document, tag: Tag):
    image = tag.find("img")
    caption = tag.find("figcaption")
    if image and image.get("src"):
        image_path = ROOT / image["src"]
        if image_path.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image_path), width=Inches(6.5))
    if caption:
        paragraph = paragraph_from_tag(document, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
        if paragraph.runs:
            paragraph.runs[0].italic = True


def note_from_tag(document: Document, tag: Tag):
    paragraph = document.add_paragraph()
    for child in tag.children:
        append_inline(paragraph, child)


def div_from_tag(document: Document, tag: Tag):
    classes = set(tag.get("class", []))
    if "page-break" in classes:
        document.add_page_break()
        return
    if "cover-meta" in classes:
        for child in tag.children:
            if isinstance(child, Tag) and child.name == "p":
                paragraph_from_tag(document, child, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    if "toc" in classes:
        for child in tag.children:
            process_tag(document, child)
        return
    if "note" in classes:
        note_from_tag(document, tag)
        return
    for child in tag.children:
        process_tag(document, child)


def process_tag(document: Document, node):
    if isinstance(node, NavigableString):
        if node.strip():
            document.add_paragraph(normalized_text(str(node)))
        return

    if not isinstance(node, Tag):
        return

    if node.name == "h1":
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(normalized_text(node.get_text(" ", strip=True)))
        run.bold = True
        run.font.size = Pt(24)
        return

    if node.name == "h2":
        document.add_heading(normalized_text(node.get_text(" ", strip=True)), level=1)
        return

    if node.name == "h3":
        document.add_heading(normalized_text(node.get_text(" ", strip=True)), level=2)
        return

    if node.name == "h4":
        document.add_heading(normalized_text(node.get_text(" ", strip=True)), level=3)
        return

    if node.name == "p":
        paragraph_from_tag(document, node)
        return

    if node.name in {"ul", "ol"}:
        list_from_tag(document, node)
        return

    if node.name == "table":
        table_from_tag(document, node)
        return

    if node.name == "figure":
        figure_from_tag(document, node)
        return

    if node.name == "div":
        div_from_tag(document, node)
        return

    for child in node.children:
        process_tag(document, child)


def build_docx():
    soup = BeautifulSoup(HTML_PATH.read_text(encoding="utf-8"), "html.parser")
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    body = soup.body
    for child in body.children:
        process_tag(document, child)

    document.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
